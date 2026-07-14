"""Document source — read plain text out of a file for optimization.

Supports Word documents (``.docx``) via ``python-docx`` and plain-text files
(``.txt``, ``.md``, ``.log``). This replaces the JIRA/DevOps connectors as the
input source: the raw text pulled here is what the optimizer shrinks before it
reaches the model.
"""

from __future__ import annotations

import os

# Extensions we treat as already-plain text.
_TEXT_EXTS = {".txt", ".md", ".markdown", ".log", ".text", ""}


def read_document(path: str) -> str:
    """Return the plain text of ``path``.

    ``.docx`` files are parsed with python-docx (paragraphs + tables). Everything
    else is read as UTF-8 text. Raises a clear error if the file is missing or a
    ``.docx`` is requested without python-docx installed.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Document not found: {path}")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        return _read_docx(path)
    if ext == ".doc":
        raise ValueError(
            "Legacy .doc files are not supported. Save the document as .docx "
            "(or .txt) and try again."
        )
    if ext in _TEXT_EXTS:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()

    # Unknown extension — best effort as text so the user isn't blocked.
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _read_docx(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - depends on optional install
        raise RuntimeError(
            "Reading .docx files requires python-docx. Install it with:\n"
            "    pip install python-docx"
        ) from exc

    document = docx.Document(path)
    parts: list[str] = [p.text for p in document.paragraphs]

    # Include table cell text too — a lot of Word content lives in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


SUPPORTED_EXTENSIONS = (".docx", ".txt", ".md", ".markdown", ".log", ".text")
